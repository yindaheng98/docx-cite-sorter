import re
from docx import Document


def parse_citation(citation: str):
    citations = []
    for c in re.findall(r"\[(.+?)]", citation):
        citations.extend(c.split(','))
    citations_num = []
    for i in range(len(citations)):
        if '-' in citations[i]:
            s, e = citations[i].split('-')[0:2]
            citations_num.extend(range(int(s), int(e) + 1))
        else:
            citations_num.append(int(citations[i]))
    return citations_num


def get_citations(path):
    document = Document(path)
    citations = []
    for par in document.paragraphs:
        citations.extend(re.findall(r"\[[0-9-,]*]", par.text))
    for i in range(len(citations)):
        citations[i] = parse_citation(citations[i])
    return citations


def get_citations_shine(citations):
    citation_shine = []
    for citation in citations:
        for c in citation:
            if c not in citation_shine:
                citation_shine.append(c)
    return citation_shine


def sort_citations(citations, citations_shine):
    citations_sorted = []
    for citation in citations:
        citation_sorted = []
        for c in citation:
            citation_sorted.append(citations_shine.index(c) + 1)
        citations_sorted.append(citation_sorted)
    return citations_sorted


def get_references(path):
    document = Document(path)
    numId = None
    references = []
    for par in document.paragraphs:
        if par._element.pPr.numPr is not None:
            num = par._element.pPr.numPr.numId.val
            if num != numId:
                references = []
                numId = num
            references.append(par)
    return references
